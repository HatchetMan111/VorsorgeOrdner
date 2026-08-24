"""DOCX-Export des Vorsorge-Ordners (python-docx)."""
from __future__ import annotations

import os
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from content import (APP_SUBTITLE, APP_TITLE, BRAND_FOOTER, CHECKLIST_ITEMS,
                     DISCLAIMER_KURZ, LEGAL_TABLE, WEGWEISER, checklist_status)
from models import VorsorgeDaten
from pdf_export import (REGISTERS, safe_filename)

GREEN_HEX = "2E7D32"
AMBER_FILL = "FFF8E1"
RED_FILL = "FDECEA"
GREEN_FILL = "E8F5E9"

BRAND = os.environ.get("VO_BRAND", "Vorsorge-Ordner")


def _shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _hint(doc: Document, text: str, title: str, fill: str):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.cell(0, 0)
    _shade(cell, fill)
    p = cell.paragraphs[0]
    run = p.add_run(f"{title}: {text}")
    run.font.size = Pt(9)
    doc.add_paragraph()


def _kv(doc: Document, label: str, value, blank="—"):
    val = str(value or "").strip()
    p = doc.add_paragraph()
    r = p.add_run(f"{label}: ")
    r.bold = True
    if val:
        p.add_run(val)
    else:
        gray = p.add_run(blank)
        gray.font.color.rgb = RGBColor(0x9E, 0x9E, 0x9E)


def _table(doc: Document, headers: list, rows: list):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        _shade(c, GREEN_HEX)
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9)
    for row in rows or [["—", "", "", ""]][:1]:
        cells = t.add_row().cells
        for i, v in enumerate(row[:len(headers)]):
            run = cells[i].paragraphs[0].add_run(str(v))
            run.font.size = Pt(9)
    doc.add_paragraph()


def _draft(doc: Document, title: str, paragraphs: list[str]):
    _hint(doc, "\n".join(p for p in paragraphs if p), title, GREEN_FILL)


def build_docx(d: VorsorgeDaten) -> bytes:
    doc = Document()
    section = doc.sections[0]
    footer_p = section.footer.paragraphs[0]
    footer_p.text = f"{BRAND_FOOTER} · erstellt am {date.today().strftime('%d.%m.%Y')}"
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Deckblatt
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run(APP_TITLE)
    r.bold = True
    r.font.size = Pt(28)
    sp = doc.add_paragraph(APP_SUBTITLE)
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    _table(doc, ["Feld", "Angabe"], [
        ("Name", d.person.name), ("Geburtsdatum", d.person.geburtsdatum),
        ("Geburtsort", d.person.geburtsort), ("Anschrift", d.person.anschrift),
        ("Vertrauensperson", d.vertrauenspersonen.haupt.name),
        ("Ersatzperson", d.vertrauenspersonen.ersatz.name),
    ])
    _hint(doc, DISCLAIMER_KURZ, "Bitte lesen", AMBER_FILL)
    doc.add_page_break()

    # Checkliste / Inhaltsverzeichnis
    doc.add_heading("Inhaltsverzeichnis & Checkliste", level=1)
    _table(doc, ["Nr", "Punkt", "Status"],
           [(str(i), titel, checklist_status(d.checkliste, i))
            for i, (titel, _) in enumerate(CHECKLIST_ITEMS, start=1)])
    doc.add_page_break()

    # Wegweiser
    doc.add_heading("Der Weg dorthin – in 8 Schritten", level=1)
    for i, (titel, desc) in enumerate(WEGWEISER, start=1):
        doc.add_heading(f"Schritt {i} · {titel}", level=3)
        doc.add_paragraph(desc)
    doc.add_page_break()

    vp = d.vertrauenspersonen

    # Register 1–21 linear
    builders_extra = {
        1: lambda: [
            _kv(doc, "Status", d.vorsorgevollmacht.status),
            _kv(doc, "Aufbewahrungsort", d.vorsorgevollmacht.aufbewahrung),
            _kv(doc, "Zentrales Vorsorgeregister", "Ja" if d.vorsorgevollmacht.register_eingetragen else "Nein"),
            _kv(doc, "Bevollmächtigte Person", d.vertrauenspersonen.haupt.name),
            _kv(doc, "Ersatz-Bevollmächtigte(r)", d.vertrauenspersonen.ersatz.name),
            _kv(doc, "Gilt über den Tod hinaus", "Ja" if d.vorsorgevollmacht.ueber_tod_hinaus else "Nein"),
            _draft(doc, "Entwurf – bitte abschreiben/anpassen und unterschreiben", [
                f"Ich, {d.person.name or '____________________'}, erteile "
                f"{d.vorsorgevollmacht.bevollmaechtigter.name or '____________________'} Vollmacht, mich in Gesundheits-, "
                "Aufenthalts-, Vermögens-, Post-, Behörden- und Bankangelegenheiten zu vertreten. "
                + ("Diese Vollmacht gilt über meinen Tod hinaus." if d.vorsorgevollmacht.ueber_tod_hinaus else ""),
                "Ort, Datum: ______________________     Unterschrift: ______________________",
            ]),
        ],
        2: lambda: [
            _kv(doc, "Status", d.patientenverfuegung.status),
            _kv(doc, "Situationen", "; ".join(d.patientenverfuegung.situationen)),
            _kv(doc, "Lebenserhaltende Maßnahmen", d.patientenverfuegung.lebenserhaltend),
            _kv(doc, "Wiederbelebung", d.patientenverfuegung.wiederbelebung),
            _kv(doc, "Künstliche Ernährung", d.patientenverfuegung.kuenstliche_ernaehrung),
            _kv(doc, "Schmerz-/Symptomlinderung", d.patientenverfuegung.schmerzlinderung),
            _draft(doc, "Entwurf", [d.patientenverfuegung.text,
                                    "Ort, Datum: ______________________     Unterschrift: ______________________"]),
        ],
        3: lambda: [
            _kv(doc, "Vorgeschlagene(r) Betreuer(in)", d.betreuungsverfuegung.betreuer.name),
            _kv(doc, "Ersatz", d.betreuungsverfuegung.ersatz_betreuer.name),
            _kv(doc, "Wunsch im Pflegefall", d.betreuungsverfuegung.pflegewunsch),
            _kv(doc, "Besondere Wünsche", d.betreuungsverfuegung.wuensche),
            _draft(doc, "Entwurf", [d.betreuungsverfuegung.wuensche]),
        ],
        4: lambda: ([] if not d.sorgerecht.relevant else [
            _kv(doc, "Zum Vormund bestimmt", d.sorgerecht.vormund.name),
            _kv(doc, "Ersatz-Vormund", d.sorgerecht.ersatz_vormund.name),
            _kv(doc, "Erziehungswünsche", d.sorgerecht.erziehungswuensche),
        ]),
        5: lambda: [
            _kv(doc, "Art", d.testament.art),
            _kv(doc, "Datum", d.testament.datum),
            _kv(doc, "Aufbewahrungsort", d.testament.aufbewahrung),
            _table(doc, ["Name", "Beziehung", "Anteil", "Ersatzerbe"],
                   [(e.name, e.beziehung, e.anteil, e.ersatzerbe) for e in d.testament.erben]),
            _kv(doc, "Vermächtnisse", d.testament.vermaechtnisse),
            _kv(doc, "Testamentsvollstrecker", d.testament.testamentsvollstrecker),
            _hint(doc, "Eigenhändiges Testament muss vollständig handschriftlich sein (§ 2247 BGB).",
                  "ACHTUNG", RED_FILL),
        ],
        6: lambda: ([] if not d.nachfolge.relevant else [
            _kv(doc, "Art der Regelung", d.nachfolge.art_regelung),
            _kv(doc, "Nachfolger(in)", d.nachfolge.nachfolger),
            _kv(doc, "Berater", d.nachfolge.berater),
            _kv(doc, "Hinweise", d.nachfolge.hinweise),
        ]),
        7: lambda: [
            _kv(doc, "Bankvollmacht erteilt", d.bank.vollmacht_erteilt),
            _kv(doc, "An wen", d.bank.vollmacht_an),
            _table(doc, ["Institut", "Ansprechpartner", "Kontoart"],
                   [(b.institut, b.ansprechpartner, b.kontoart) for b in d.bank.banken]),
        ],
        8: lambda: [_kv(doc, "Über den Tod hinaus geregelt",
                        "Ja" if d.vorsorgevollmacht.ueber_tod_hinaus else "Bitte in Register 1 festlegen")],
        9: lambda: [
            _kv(doc, "Bestattungsart", d.bestattung.art),
            _kv(doc, "Friedhof/Ort", d.bestattung.friedhof),
            _kv(doc, "Grabart", d.bestattung.grabart),
            _kv(doc, "Trauerfeier", d.bestattung.trauerfeier),
            _kv(doc, "Musikwünsche", d.bestattung.musik),
            _kv(doc, "Blumen/Spende", d.bestattung.blumen_spende),
            _kv(doc, "Traueranzeige", d.bestattung.anzeige),
        ],
        10: lambda: [_kv(doc, "Entscheidung", d.organspende.entscheidung),
                     _kv(doc, "Details", d.organspende.details)],
        11: lambda: [
            _kv(doc, "Passwort-Manager", d.digital.passwort_manager),
            _kv(doc, "Master-Zugang (handschriftlich ergänzen)", ""),
            _table(doc, ["Dienst", "Benutzername", "Aktion"],
                   [(k.dienst, k.benutzername, k.aktion) for k in d.digital.konten]),
        ],
        12: lambda: [_kv(doc, "Schlüssel liegen bei", d.schluessel),
                     _kv(doc, "Zugänge/Codes", d.zugaenge)],
        13: lambda: [_table(doc, ["Rolle", "Name", "Telefon"],
                            [(k.rolle, k.name, k.telefon) for k in d.kontakte])],
        14: lambda: [_table(doc, ["Art", "Partner", "Kündigungsfrist", "Ort"],
                            [(v.art, v.partner, v.kuendigungsfrist, v.ort_unterlagen) for v in d.vertrage])],
        15: lambda: [
            _kv(doc, "Konten/Depots", d.vermoegen.konten),
            _kv(doc, "Immobilien", d.vermoegen.immobilien),
            _kv(doc, "Fahrzeuge", d.vermoegen.fahrzeuge),
            _kv(doc, "Schulden/Kredite", d.vermoegen.schulden_kredite),
            _kv(doc, "Wertgegenstände", d.vermoegen.wertgegenstaende),
        ],
        16: lambda: [_table(doc, ["Art", "Gesellschaft", "Standort"],
                            [(v.art, v.gesellschaft, v.police_ort) for v in d.versicherungen])],
        17: lambda: [_kv(doc, "Rente", d.rente_steuer.rente),
                     _kv(doc, "Steuer", d.rente_steuer.steuer),
                     _kv(doc, "Ort der Unterlagen", d.rente_steuer.unterlagen_ort)],
        18: lambda: [
            _kv(doc, "Geburtsurkunde", d.urkunden.geburtsurkunde),
            _kv(doc, "Heiratsurkunde", d.urkunden.heiratsurkunde),
            _kv(doc, "Scheidungsurteil", d.urkunden.scheidungsurteil),
            _kv(doc, "Ausweis-Kopie", d.urkunden.ausweiskopie),
        ],
        19: lambda: [
            _kv(doc, "Blutgruppe", d.medizin.blutgruppe),
            _kv(doc, "Hausarzt", d.medizin.hausarzt_name),
            _kv(doc, "Telefon Hausarzt", d.medizin.hausarzt_telefon),
            _table(doc, ["Arzt", "Fachrichtung", "Telefon"],
                   [(a.name, a.fachrichtung, a.telefon) for a in d.medizin.aerzte]),
            _hint(doc, d.medizin.allergien or "— keine Angaben —", "ALLERGIEN", RED_FILL),
            _kv(doc, "Diagnosen", d.medizin.diagnosen),
            _table(doc, ["Medikament", "Dosierung"],
                   [(m.name, m.dosierung) for m in d.medizin.medikamente]),
            _kv(doc, "Medikamentenplan-Ort", d.medizin.medikamentenplan_ort),
        ],
        20: lambda: [_table(doc, ["Name", "Tierart", "Betreuungsperson", "Tierarzt"],
                            [(t_.name, t_.tierart, t_.betreuungsperson, t_.tierarzt) for t_ in d.haustiere])],
        21: lambda: [_kv(doc, "Dokumente liegen bei", d.notfallkarte.aufbewahrung),
                     _kv(doc, "Kontaktperson", d.notfallkarte.kontakt),
                     _kv(doc, "Telefon", d.notfallkarte.telefon)],
    }

    hints = {
        1: CHECKLIST_ITEMS[0][1], 2: CHECKLIST_ITEMS[1][1], 3: CHECKLIST_ITEMS[2][1],
        4: CHECKLIST_ITEMS[3][1], 5: CHECKLIST_ITEMS[4][1], 6: CHECKLIST_ITEMS[5][1],
        7: CHECKLIST_ITEMS[6][1], 8: CHECKLIST_ITEMS[7][1], 9: CHECKLIST_ITEMS[8][1],
        10: CHECKLIST_ITEMS[9][1], 11: CHECKLIST_ITEMS[10][1],
    }
    for nr, (titel, _fn) in enumerate(REGISTERS, start=1):
        doc.add_page_break()
        h = doc.add_heading("", level=1)
        run = h.add_run(f"Register {nr} · {titel}")
        run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
        if nr in hints:
            _hint(doc, hints[nr], "Rechtlicher Hinweis", AMBER_FILL)
        for flowable in builders_extra.get(nr, lambda: [])():
            pass

    # Rechtliche Hinweise
    doc.add_page_break()
    doc.add_heading("Rechtliche Hinweise & Formvorschriften", level=1)
    _table(doc, ["Dokument", "Formvorschrift", "Notar nötig?"], LEGAL_TABLE)
    for txt in [
        "Lass Vollmachten und Testament von Notar bzw. Anwalt prüfen.",
        "Sprich deine Patientenverfügung mit deinem Hausarzt durch.",
        "Registriere die Vorsorgevollmacht unter vorsorgeregister.de.",
        "Gib Kopien an deine Vertrauenspersonen und prüfe alles jährlich.",
    ]:
        doc.add_paragraph(txt, style="List Bullet")
    _hint(doc, DISCLAIMER_KURZ + " Es werden keine Rechts-, Steuer- oder Elektroleistungen erbracht.",
          "Hinweis", AMBER_FILL)

    import io
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
